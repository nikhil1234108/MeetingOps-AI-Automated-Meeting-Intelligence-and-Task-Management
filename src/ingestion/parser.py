import os
import re
import json
from typing import Dict, Any, List, Tuple
from pypdf import PdfReader
from docx import Document

class TranscriptParser:
    """
    Standardizes various input meeting notes or transcripts (VTT, SRT, PDF, DOCX, JSON, TXT)
    into a uniform transcript structure: [Timestamp] Speaker Name: Statement.
    """

    @staticmethod
    def clean_text(text: str) -> str:
        """Removes duplicate whitespace, empty lines, and corrupted characters."""
        # Split into lines and strip each line of trailing/leading spaces
        lines = [line.strip() for line in text.splitlines()]
        # Remove all empty lines
        cleaned_lines = []
        for line in lines:
            if line:
                # Replace multiple spaces inside lines with a single space
                line = re.sub(r'[ \t]+', ' ', line)
                cleaned_lines.append(line)
            
        return "\n".join(cleaned_lines)

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Parses a standard text file or markdown file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return TranscriptParser.clean_text(content)

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Parses a PDF file page by page using pypdf."""
        reader = PdfReader(file_path)
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages_text.append(text)
        
        full_text = "\n".join(pages_text)
        if not full_text.strip():
            raise ValueError("PDF text extraction returned empty string. PDF might be scanned and require OCR.")
        return TranscriptParser.clean_text(full_text)

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Parses a Microsoft Word Document (.docx)."""
        doc = Document(file_path)
        paragraphs_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs_text.append(para.text)
        
        # Also parse tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs_text.append(" | ".join(row_text))

        full_text = "\n".join(paragraphs_text)
        return TranscriptParser.clean_text(full_text)

    @staticmethod
    def parse_json(file_path: str) -> str:
        """
        Parses a structured JSON meeting record.
        Expects a list of turns/messages or a nested structure.
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # If it's a list, format it
        turns = []
        if isinstance(data, list):
            for i, item in enumerate(data):
                speaker = item.get("speaker") or item.get("user") or item.get("name") or f"Speaker {i+1}"
                timestamp = item.get("timestamp") or item.get("time") or ""
                text = item.get("text") or item.get("message") or item.get("statement") or ""
                
                time_str = f"[{timestamp}] " if timestamp else ""
                turns.append(f"{time_str}{speaker}:\n{text}")
        elif isinstance(data, dict):
            # Check for a 'transcript' or 'messages' key
            messages = data.get("messages") or data.get("turns") or data.get("transcript")
            if isinstance(messages, list):
                for i, item in enumerate(messages):
                    speaker = item.get("speaker") or item.get("user") or item.get("name") or f"Speaker {i+1}"
                    timestamp = item.get("timestamp") or item.get("time") or ""
                    text = item.get("text") or item.get("message") or item.get("statement") or ""
                    
                    time_str = f"[{timestamp}] " if timestamp else ""
                    turns.append(f"{time_str}{speaker}:\n{text}")
            else:
                # Direct key-value conversion fallback
                return TranscriptParser.clean_text(json.dumps(data, indent=2))
        else:
            return TranscriptParser.clean_text(str(data))
            
        return "\n\n".join(turns)

    @staticmethod
    def parse_vtt(file_path: str) -> str:
        """
        Parses a WebVTT (.vtt) file.
        Removes headers, cue timing numbers, and HTML-like speaker tags,
        leaving clean speaker dialogue lines with timestamps.
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()

        transcript_lines = []
        current_time = ""
        
        time_regex = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d{3}|\d{2}:\d{2}\.\d{3})\s*-->')
        
        # WebVTT tag pattern: <v Speaker Name>Speaker dialogue</v>
        webvtt_v_tag = re.compile(r'<v\s+([^>]+)>(.*?)</v>', re.DOTALL)
        # General tag pattern to strip residual tags
        html_tag = re.compile(r'<[^>]+>')

        for line in lines:
            line_str = line.strip()
            
            # Skip WebVTT signature header
            if "WEBVTT" in line_str:
                continue
            
            # Match timestamp
            match = time_regex.search(line_str)
            if match:
                # Keep HH:MM:SS part (first 8 chars if HH:MM:SS, else MM:SS)
                full_time = match.group(1)
                current_time = full_time.split('.')[0] # Strip milliseconds
                continue
            
            if not line_str or line_str.isdigit():
                continue
            
            # Check for <v Speaker> tag
            v_match = webvtt_v_tag.search(line_str)
            if v_match:
                speaker = v_match.group(1).strip()
                statement = v_match.group(2).strip()
                statement = html_tag.sub('', statement) # clean internal tags
                time_prefix = f"[{current_time}] " if current_time else ""
                transcript_lines.append(f"{time_prefix}{speaker}:\n{statement}")
            else:
                # Fallback to general lines: check if there's an implicit speaker (e.g., "Nancy Collins:")
                cleaned_line = html_tag.sub('', line_str)
                if ":" in cleaned_line:
                    parts = cleaned_line.split(":", 1)
                    speaker = parts[0].strip()
                    statement = parts[1].strip()
                    time_prefix = f"[{current_time}] " if current_time else ""
                    transcript_lines.append(f"{time_prefix}{speaker}:\n{statement}")
                else:
                    # Append dialogue line to previous turn if no speaker
                    if transcript_lines:
                        transcript_lines[-1] += f" {cleaned_line}"
                    else:
                        time_prefix = f"[{current_time}] " if current_time else ""
                        transcript_lines.append(f"{time_prefix}Unknown Speaker:\n{cleaned_line}")

        return "\n\n".join(transcript_lines)

    @staticmethod
    def parse_srt(file_path: str) -> str:
        """
        Parses an SRT subtitle file.
        Format is:
        1
        00:01:20,000 --> 00:01:23,000
        Speaker Name: Dialogue line
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()

        transcript_lines = []
        current_time = ""
        
        time_regex = re.compile(r'(\d{2}:\d{2}:\d{2})[,\.]\d{3}\s*-->')
        
        for line in lines:
            line_str = line.strip()
            
            if not line_str or line_str.isdigit():
                continue
            
            match = time_regex.search(line_str)
            if match:
                current_time = match.group(1)
                continue
            
            # Dialogue block
            if ":" in line_str:
                parts = line_str.split(":", 1)
                speaker = parts[0].strip()
                statement = parts[1].strip()
                time_prefix = f"[{current_time}] " if current_time else ""
                transcript_lines.append(f"{time_prefix}{speaker}:\n{statement}")
            else:
                if transcript_lines:
                    transcript_lines[-1] += f" {line_str}"
                else:
                    time_prefix = f"[{current_time}] " if current_time else ""
                    transcript_lines.append(f"{time_prefix}Unknown Speaker:\n{line_str}")
                    
        return "\n\n".join(transcript_lines)

    @classmethod
    def parse(cls, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Ingests the input file, identifies type, normalizes contents,
        runs pre-AI validation, and returns (standardized_text, metadata_report).
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Input file not found at: {file_path}")

        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower()
        file_size = os.path.getsize(file_path)

        # Ingestion routing
        if file_ext in ('.txt', '.md'):
            normalized_text = cls.parse_txt(file_path)
        elif file_ext == '.pdf':
            normalized_text = cls.parse_pdf(file_path)
        elif file_ext == '.docx':
            normalized_text = cls.parse_docx(file_path)
        elif file_ext == '.json':
            normalized_text = cls.parse_json(file_path)
        elif file_ext == '.vtt':
            normalized_text = cls.parse_vtt(file_path)
        elif file_ext == '.srt':
            normalized_text = cls.parse_srt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Please provide TXT, PDF, DOCX, JSON, VTT, or SRT.")

        # Pre-AI Validation Checks
        if not normalized_text.strip():
            raise ValueError(f"File ingestion resulted in an empty transcript for {file_name}.")
        
        if len(normalized_text) < 50:
            raise ValueError("Transcript content is too short (minimum threshold is 50 characters).")

        # Detect participants
        # We search for "[Timestamp]\nSpeaker Name:\n" or "Speaker Name:" pattern
        speaker_matches = re.findall(r'(?:^|\n)(?:\[\d{2}:\d{2}:\d{2}\]\s*)?([^:\n]+):', normalized_text)
        unique_speakers = sorted(list(set([s.strip() for s in speaker_matches if s.strip() and s.strip() != "Unknown Speaker"])))
        
        if not unique_speakers:
            unique_speakers = ["Unknown Speaker"]

        # Estimate meeting duration from WebVTT/SRT timestamps if available
        # Find first and last timestamps
        timestamps = re.findall(r'\[(\d{2}:\d{2}:\d{2})\]', normalized_text)
        duration_est = "UNKNOWN"
        if len(timestamps) >= 2:
            try:
                def to_seconds(t_str):
                    parts = list(map(int, t_str.split(':')))
                    return parts[0]*3600 + parts[1]*60 + parts[2]
                sec_diff = to_seconds(timestamps[-1]) - to_seconds(timestamps[0])
                duration_est = f"{sec_diff // 60} minutes, {sec_diff % 60} seconds"
            except Exception:
                pass

        # Package standardized format
        standardized_transcript = (
            f"Meeting Information\n"
            f"File Name: {file_name}\n"
            f"Detected Participants: {', '.join(unique_speakers)}\n"
            f"Estimated Duration: {duration_est}\n\n"
            f"Transcript:\n{normalized_text}"
        )

        metadata_report = {
            "file_name": file_name,
            "file_type": file_ext[1:].upper(),
            "file_size_bytes": file_size,
            "encoding": "UTF-8",
            "detected_participants": unique_speakers,
            "participant_count": len(unique_speakers),
            "estimated_duration": duration_est
        }

        return standardized_transcript, metadata_report
